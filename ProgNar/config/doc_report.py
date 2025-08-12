import json
from heapq import merge

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from tkinter import filedialog
from datetime import datetime
from config.utils import resource_path
from config.client_utils import get_client_data



def remove_cell_borders(cell):
    """Usuwa obramowanie komórki."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for border in ['top', 'left', 'bottom', 'right']:
        border_elem = tcPr.first_child_found_in(f'w:{border}')
        if border_elem is not None:
            border_elem.set(qn('w:val'), 'none')

def remove_all_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)

    for position in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        border = borders.find(qn(f'w:{position}'))
        if border is None:
            border = OxmlElement(f'w:{position}')
            borders.append(border)
        border.set(qn('w:val'), 'nil')

def remove_cell_border(cell, border_position):
    """Usuwa WYBRANE obramowanie WYBRANEJ komórki."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)

    border = borders.find(qn(f'w:{border_position}'))
    if border is None:
        border = OxmlElement(f'w:{border_position}')
        borders.append(border)

    border.set(qn('w:val'), 'nil')  # Usuwa obramowanie

def extract_price(cell):
    try:
        text = cell.text.strip().replace(" zł", "").replace(",", ".")
        return float(text)
    except:
        return 0.0


def generate_report(client_name, main_app):
    json_path = resource_path("data/temp_cart.json")
    with open(json_path, encoding='utf-8') as f:
        data = json.load(f)

    # Pobierz dane klienta za pomocą get_client_data
    client_data = get_client_data(client_name.get(), main_app)

    # Przygotuj dane klienta, używając "......." dla pustych pól
    client_name_str = client_name.get().strip() if client_name.get().strip() and client_name.get() != "- -" else "......."
    client_address = client_data.get("address", "").strip() if client_data else "......."
    client_contact = client_data.get("contact", "").strip() if client_data else "......."

    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Dokument Word", "*.docx")],
        title="Zapisz raport jako"
    )
    if not save_path:
        print("Zapis anulowany.")
        return

    def write_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        paragraph.alignment = align

    def count_pluses(table, threshold=12, below=True):
        """
        Zlicza plusy w kolumnie ciecia, gdzie wartość w kolumnie 2 (srednica) spełnia warunek względem progu.

        :param table: tabela (np. z python-docx)
        :param threshold: wartość progowa (domyślnie 12)
        :param below: jeśli True, liczy plusy z wartością < threshold; jeśli False, >= threshold
        :return: liczba spełniających warunek plusów
        """
        total_qty = 0
        for row in table.rows:
            cells = row.cells
            if len(cells) > 7:
                if cells[6].text.strip() == "+":
                    diameter = float(cells[2].text.strip())
                    if (below and diameter < threshold) or (not below and diameter >= threshold):
                        try:
                            qty = int(cells[7].text.strip())
                            total_qty += qty
                        except ValueError:
                            pass  # Jeśli nie da się sparsować liczby, ignorujemy
        return total_qty

    doc = Document()

    # Ustaw minimalne marginesy
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    total_width_cm = 22.2 - 0.6 * 2  # A4 szerokość - marginesy boczne

    # Czcionka
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(14)  # Domyślna czcionka dla dokumentu

    # Nagłówek "WYCENA"
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("          WYCENA           ")
    header_run.bold = True
    header_run.underline = True
    header_run.font.size = Pt(16)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(10)

    # Dodajemy dzisiejszą datę przed nazwą firmy
    today = datetime.today().strftime("%d.%m.%Y")
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Świdnica, {today}      ")
    date_run.font.size = Pt(10)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.paragraph_format.space_after = Pt(8)  # Odstęp po dacie

    # Górny wiersz z nazwą klienta, adresem, kontaktem i logo
    top_table = doc.add_table(rows=1, cols=2)

    # Komórka z nazwą klienta, adresem i kontaktem
    client_cell = top_table.cell(0, 0)
    client_para = client_cell.paragraphs[0]
    client_run = client_para.add_run(client_name_str)
    client_run.bold = True
    client_run.font.size = Pt(14)
    client_para.paragraph_format.space_after = Pt(2)  # Minimalny odstęp po nazwie

    # Dodanie linii z adresem i kontaktem
    address_para = client_cell.add_paragraph(client_address)
    address_para.runs[0].font.size = Pt(14)
    address_para.paragraph_format.space_after = Pt(2)  # Minimalny odstęp po adresie

    contact_para = client_cell.add_paragraph(client_contact)
    contact_para.runs[0].font.size = Pt(14)
    contact_para.paragraph_format.space_after = Pt(2)  # Minimalny odstęp po kontakcie

    logo_cell = top_table.cell(0, 1)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        logo_para.add_run().add_picture(resource_path("img/logo.png"), width=Inches(2))
    except:
        logo_para.add_run("Brak logo")

    doc.add_paragraph()

    # Tabela danych
    headers = [
        "L.P.", "Typ", "Ø OD", "Ø Chw", "L[mm]",
        "Ilość ostrzy", "Cięcie", "Ilość szt.",
        "Cena ostrzenia", "Wartość ostrzenia",
        "Powłoka", "Cena powłoki", "Wartość powłoki", "Uwagi"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False

    # Wyznaczenie szerokości kolumn
    fixed_widths_cm = {
        0: 0.8,
        1: 1.8,
        2: 1.1,
        3: 1.1,
        4: 1.1,
        5: 1.2,
        6: 1.1,
        7: 1.1
    }
    remaining_columns = list(range(8, 14))  # pozostałe
    used_width = sum(fixed_widths_cm.values())
    remaining_width = total_width_cm - used_width
    equal_width = remaining_width / len(remaining_columns)

    column_widths_cm = {}
    column_widths_cm.update(fixed_widths_cm)
    for i in remaining_columns:
        column_widths_cm[i] = equal_width

    # Nagłówki
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(7)  # Czcionka nagłówków tabeli na 7 pt
        if i < 7:  # Usuwanie obramowań dla kolumn 0–6
            remove_cell_borders(cell)

    # Dane
    sum_qty = 0
    sum_sharpening = 0
    sum_coating = 0

    for i, item in enumerate(data['items'], start=1):
        p = item['params']
        qty = item['quantity']
        sharpen = item['sharpening_price']
        coat = item['coating_price']

        val_sharpen = sharpen * qty
        val_coat = coat * qty

        sum_qty += qty
        sum_sharpening += val_sharpen
        sum_coating += val_coat

        row = table.add_row().cells
        values = [
            str(i), p.get("Typ", "-"), p.get("Srednica", "-"),
            p.get("fiChwyt", "-"), p.get("Długość całkowita", "-"),
            p.get("Ilosc ostrzy", "-"), p.get("ciecie", "-"), str(qty),
            f"{sharpen:.2f} PLN", f"{val_sharpen:.2f} PLN",
            p.get("Powloka", "-"), f"{coat:.2f} PLN" if coat else "-",
            f"{val_coat:.2f} PLN" if coat else "-", p.get("Uwagi", "-")
        ]

        for col, val in enumerate(values):
            cell = row[col]
            para = cell.paragraphs[0]
            para.text = val
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].font.size = Pt(7)  # Czcionka danych tabeli na 7 pt
            if col < 7:  # Usuwanie obramowań dla kolumn 0–6
                remove_cell_borders(cell)

    # Pusty wiersz zmergowany
    empty_row = table.add_row().cells
    empty_row[0].merge(empty_row[13])  # Mergowanie wszystkich komórek (0–13)
    remove_cell_borders(empty_row[0])  # Usunięcie obramowań dla połączonej komórki

    # Wiersz podsumowania
    summary_row = table.add_row().cells
    summary_row[0].merge(summary_row[6])  # Mergowanie komórek 0–6

    # Komórka 0
    paragraph = summary_row[0].paragraphs[0]
    paragraph.clear()  # Czyści istniejący tekst
    run = paragraph.add_run("Liczba sztuk:")
    run.bold = True
    run.font.size = Pt(7)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Komórka 7
    paragraph = summary_row[7].paragraphs[0]
    paragraph.clear()  # Czyści istniejący tekst
    run = paragraph.add_run(str(sum_qty))
    run.bold = True
    run.font.size = Pt(7)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Komórka 8
    paragraph = summary_row[8].paragraphs[0]
    paragraph.clear()  # Czyści istniejący tekst
    run = paragraph.add_run("Suma:")
    run.bold = True
    run.font.size = Pt(7)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Komórka 9
    paragraph = summary_row[9].paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(f"{sum_sharpening:.2f} PLN")
    run.bold = True
    run.font.size = Pt(7)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Komórka 11
    paragraph = summary_row[11].paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run("Suma:")
    run.bold = True
    run.font.size = Pt(7)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Komórka 12
    paragraph = summary_row[12].paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(f"{sum_coating:.2f} PLN")
    run.bold = True
    run.font.size = Pt(7)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    remove_cell_border(summary_row[11], 'right')  # Usuwa prawą krawędź komórki 11
    remove_cell_border(summary_row[12], 'left')  # Usuwa lewą krawędź komórki 12
    remove_cell_border(summary_row[8], 'right')  # Usuwa prawą krawędź komórki 11
    remove_cell_border(summary_row[9], 'left')  # Usuwa lewą krawędź komórki 12
    remove_cell_border(summary_row[0], 'right')  # Usuwa prawą krawędź komórki 11
    remove_cell_border(summary_row[7], 'left')  # Usuwa lewą krawędź komórki 12

    #wiersze uslug
    #CIECIE DO 12
    cutting_row_under = table.add_row().cells
    cutting_row_under[1].merge(cutting_row_under[2])
    cutting_row_under[4].merge(cutting_row_under[5])
    cutting_row_under[6].merge(cutting_row_under[13])
    write_cell(cutting_row_under[1],"Cięcie do D12: ",bold = True, size=8,align=WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(cutting_row_under[3], f"{count_pluses(table,12,True)} szt.", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(cutting_row_under[4], f"{count_pluses(table,12,True) * 12:.2f} zł", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    #CIECIE PONAD 12
    cutting_row_over = table.add_row().cells
    cutting_row_over[1].merge(cutting_row_over[2])
    cutting_row_over[4].merge(cutting_row_over[5])
    cutting_row_over[6].merge(cutting_row_over[13])
    write_cell(cutting_row_over[1], "Cięcie ponad D12: ", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(cutting_row_over[3], f"{count_pluses(table, 12, False)} szt.", bold=True, size=8,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(cutting_row_over[4], f"{count_pluses(table, 12, False) * 16:.2f} zł", bold=True, size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    #SZYJKA
    reduction = table.add_row().cells
    reduction[1].merge(reduction[2])
    reduction[4].merge(reduction[5])
    reduction[6].merge(reduction[13])
    write_cell(reduction[1], "Zaniżenie:", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(reduction[3], f"TBD szt.", bold=True, size=8,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(reduction[4], f"{0:.2f} zł", bold=True, size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    # Suma
    suma_uslugi = (
    extract_price(cutting_row_under[4]) +
    extract_price(cutting_row_over[4]) +
    extract_price(reduction[4])
)
    attachments_summary = table.add_row().cells
    attachments_summary[1].merge(attachments_summary[3])
    attachments_summary[4].merge(attachments_summary[5])
    attachments_summary[6].merge(attachments_summary[13])
    write_cell(attachments_summary[1], "Suma:", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(attachments_summary[4], f"{suma_uslugi:.2f} zł", bold=True, size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER)


    #koncowy wiersz
    last_row = table.add_row().cells
    last_row[0].merge(last_row[11])  # Mergowanie komórek 0–11
    last_row[12].merge(last_row[13])  # Mergowanie komórek 12–13

    write_cell(last_row[0], "Wartość zamówienia:", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write_cell(last_row[12], f"{sum_sharpening + sum_coating + extract_price(attachments_summary[4]):.2f} PLN", bold=True, size=10,
               align=WD_ALIGN_PARAGRAPH.LEFT)

    remove_all_borders(cutting_row_under[6])
    remove_all_borders(cutting_row_over[6])
    remove_all_borders(reduction[6])

    # Ustawienia szerokości kolumn
    for i, col in enumerate(table.columns):
        width_cm = column_widths_cm.get(i, 1.5)  # Domyślna szerokość 1.5 cm, jeśli nie określono
        try:
            col.width = Inches(width_cm / 2.54)  # Ustaw szerokość kolumny
        except Exception as e:
            print(f"Błąd przy ustawianiu szerokości kolumny {i}: {str(e)}")

    doc.save(save_path)
    print(f"Raport zapisany: {save_path}")